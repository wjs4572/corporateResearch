# Company Report Generator
# -----------------------
# Generates a formatted DOCX report for a company using OpenAI API and Markdown-style formatting.
# Handles logging, environment variable loading, and command-line argument parsing.

import openai
import docx
import os
import argparse
from dotenv import load_dotenv
import datetime
import inspect

def log_to_file(message):
    """
    Logs a message to a daily log file in the 'logs' directory, including filename and timestamp.
    Also prints the message to the console.
    """
    logs_dir = os.path.join(os.path.dirname(__file__), "logs")
    os.makedirs(logs_dir, exist_ok=True)
    today = datetime.datetime.now().strftime("%Y-%m-%d")
    log_file = os.path.join(logs_dir, f"report_log_{today}.txt")
    print("Logging to:", log_file)
    frame = inspect.currentframe().f_back
    filename = os.path.basename(frame.f_code.co_filename)
    timestamp = datetime.datetime.now().strftime("%H:%M:%S")
    log_entry = f"{filename} {timestamp} — {message}"
    with open(log_file, "a", encoding="utf-8") as f:
        f.write(log_entry + "\n")

# Environment Variable Loading
load_dotenv()
OPENAI_API_KEY = os.getenv('OPENAI_API_KEY')

def generate_docx_report(prompt, substitutions, model="gpt-3.5-turbo"):
    """
    Generates a DOCX report for a company using a prompt and substitution variables.
    - Substitutes <KEY> in the prompt with provided values.
    - Calls OpenAI API to generate report content.
    - Formats response with Markdown-style headings, bold, lists, and page breaks.
    - Saves the report and logs actions.
    Returns the output file path or None on error.
    """
    try:
        # Substitute variables in the prompt (replace <KEY> with value)
        for key, value in substitutions.items():
            prompt = prompt.replace(f"<{key}>", value)

        # Prepare output folder and file path first
        out_dir = os.path.join(os.path.dirname(__file__), "out")
        os.makedirs(out_dir, exist_ok=True)
        company_name = substitutions.get("CORPORATE_NAME", "company").replace(" ", "_").replace("/", "-")
        filename = f"company_report_{company_name}.docx"
        out_path = os.path.join(out_dir, filename)
        log_message = ""
        if os.path.exists(out_path):
            log_message = f"Report already exists: {out_path}. Skipping generation."
            print(log_message)
            log_to_file(log_message)
            return out_path

        # Use new OpenAI API (v1.x)
        client = openai.OpenAI(api_key=OPENAI_API_KEY)
        response = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": prompt}]
        )

        # Get the response text
        result_text = response.choices[0].message.content

        # Create a new Word document with formatted response
        doc = docx.Document()
        doc.add_heading('Company Report', 0)

        lines = result_text.splitlines()
        in_list = False
        for line in lines:
            line = line.strip()
            if not line:
                in_list = False
                continue
            # Markdown-style headings: #, ##, ###
            if line.startswith('### '):
                heading_text = line[4:].strip()
                if heading_text.startswith('Appendix') or heading_text.startswith('References'):
                    doc.add_page_break()
                doc.add_heading(heading_text, level=3)
                in_list = False
            elif line.startswith('## '):
                heading_text = line[3:].strip()
                if heading_text.startswith('Appendix') or heading_text.startswith('References'):
                    doc.add_page_break()
                doc.add_heading(heading_text, level=2)
                in_list = False
            elif line.startswith('# '):
                heading_text = line[2:].strip()
                if heading_text.startswith('Appendix') or heading_text.startswith('References'):
                    doc.add_page_break()
                doc.add_heading(heading_text, level=1)
                in_list = False
            # Bold text: **text** anywhere in the line
            elif '**' in line:
                # Split line by '**' and alternate bold/non-bold
                parts = line.split('**')
                p = doc.add_paragraph()
                for i, part in enumerate(parts):
                    if i % 2 == 1:
                        run = p.add_run(part)
                        run.bold = True
                    else:
                        p.add_run(part)
                in_list = False
            # Bulleted list: lines starting with -
            elif line.startswith('- '):
                if not in_list:
                    in_list = True
                doc.add_paragraph(line[2:], style='List Bullet')
            # Regular paragraph
            else:
                doc.add_paragraph(line)

        doc.save(out_path)
        log_message = f"Report saved to {out_path}"
        print(f"{log_message}")
        log_to_file(log_message)
        return out_path
    except Exception as exc:
        error_message = f"Error: {str(exc)}"
        print(f"{error_message}")
        log_to_file(error_message)
        return None

if __name__ == "__main__":
    # Parses command-line arguments and calls the report generation function.
    parser = argparse.ArgumentParser(description="Generate a company report using OpenAI and a template.")
    parser.add_argument('--prompt', type=str, required=True, help='Prompt/template to send to OpenAI (use <KEY> for substitutions)')
    parser.add_argument('--sub', nargs='*', help='Substitution key=value pairs')
    parser.add_argument('--model', type=str, default='gpt-3.5-turbo', help='OpenAI model to use')
    args = parser.parse_args()

    substitutions = {}
    if args.sub:
        for pair in args.sub:
            if '=' in pair:
                key, value = pair.split('=', 1)
                substitutions[key] = value

    generate_docx_report(args.prompt, substitutions, model=args.model)
    